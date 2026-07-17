"""python-docx 用の安全ヘルパー集（macOS Pages 互換）。

Markdown → docx 変換のビルダースクリプトから import して使う。
Pages で誤レイアウトを起こす OOXML 機能（keepNext / pageBreakBefore プロパティ）を
使わず、実レンダリングで検証済みのパターンだけを提供する。

使い方の詳細・禁止事項 → スキルの references/pages-compat.md
"""

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 文書ごとに上書きしてよいモジュール既定値
FONT = "游ゴシック"
HEADER_BG = "0090D0"   # 表ヘッダー行の背景色
LABEL_BG = "E6F4FA"    # ラベルセルの背景色


def setup_a4(doc, margin_cm=2.0):
    """A4 縦 + マージンを明示する。

    python-docx の既定は US Letter。A4 前提で表幅を設計すると
    ページと食い違い、Pages でのページ割り計算が狂う。
    """
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return 21.0 - margin_cm * 2  # コンテンツ幅 (cm)。表の列幅設計に使う


def set_font(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)


def shade_cell(cell, color):
    """セル背景色。必ず set_cell_text より先に呼ぶ（w:shd は w:vAlign より
    前に置く必要があり、逆順だと OOXML スキーマ違反になる）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def set_cell_text(cell, text, size=10, bold=False, color=None, align=None,
                  vcenter=True, bg=None):
    """セルにテキストを設定する。背景色は bg で渡せば正しい順序で適用される。"""
    if bg:
        shade_cell(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    if vcenter:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_row_height(row, cm_height):
    """行の最小高さ。python-docx の API のみ使う（trHeight の手動 append は
    要素が二重になりレンダラーを混乱させる）。"""
    row.height = Cm(cm_height)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def prevent_row_split(row):
    """1 行がページ境界で上下に裂けるのを防ぐ（w:cantSplit）。"""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def repeat_as_header_row(row):
    """表が複数ページにまたがるとき、この行を各ページ先頭に再表示する
    （w:tblHeader）。ヘッダー行にのみ使う。"""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))


def add_heading(doc, text, size=14, page_break_before=False,
                space_before=16, space_after=8):
    """見出し段落。keep_with_next は絶対に使わない（Pages が見出し以降の
    ブロックを丸ごと次ページへ送る）。改ページは w:br 方式で行う。"""
    if page_break_before:
        add_page_break(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    return p


def add_page_break(doc):
    """改ページ。paragraph_format.page_break_before は使わない（直後に表が
    あると Pages が表だけをさらに次ページへ送る）。改ページ文字
    （w:br type="page"）の独立段落として挿入する。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.font.size = Pt(2)  # 改ページ用の空段落が前ページの高さを消費しないよう最小化
    run.add_break(WD_BREAK.PAGE)


def add_body(doc, text, size=11, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def add_bullet(doc, text, size=11, indent_cm=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(f"・{text}")
    set_font(run, size=size)
    return p


def make_table(doc, headers, col_widths_cm, rows, row_height_cm=None,
               body_size=10, header_bg=None, repeat_header=True):
    """ヘッダー付きの表。

    - table.columns[i].width を必ず設定する（tblGrid とセル幅 tcW が
      一致しないと Pages が表を丸ごと次ページへ送る）
    - col_widths_cm の合計はコンテンツ幅（setup_a4 の戻り値）に収めること
    - 全行 cantSplit、ヘッダー行は tblHeader（repeat_header=False で無効化）
    """
    header_bg = header_bg or HEADER_BG
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, w in enumerate(col_widths_cm):
        table.columns[i].width = Cm(w)
    header_row = table.rows[0]
    prevent_row_split(header_row)
    if repeat_header:
        repeat_as_header_row(header_row)
    for i, h in enumerate(headers):
        header_row.cells[i].width = Cm(col_widths_cm[i])
        set_cell_text(header_row.cells[i], h, size=body_size, bold=True,
                      color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER,
                      bg=header_bg)
    for r, row_vals in enumerate(rows, start=1):
        row = table.rows[r]
        if row_height_cm:
            set_row_height(row, row_height_cm)
        prevent_row_split(row)
        for i, w in enumerate(col_widths_cm):
            row.cells[i].width = Cm(w)
        for i, val in enumerate(row_vals):
            set_cell_text(row.cells[i], val, size=body_size)
    return table


def finalize(doc, output_path):
    """w:zoom の percent 属性を補って保存する（属性が無いと docx スキーマ
    検証に落ちる）。"""
    settings_el = doc.settings.element
    zoom = settings_el.find(qn("w:zoom"))
    if zoom is None:
        zoom = OxmlElement("w:zoom")
        settings_el.insert(0, zoom)
    zoom.set(qn("w:percent"), "100")
    doc.save(output_path)
    print(f"written: {output_path}")
